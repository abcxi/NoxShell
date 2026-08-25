from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "output"
MOCKUP_PATH = OUTPUT_DIR / "玄壳-桌面端效果图-v2.png"
DOCX_PATH = OUTPUT_DIR / "玄壳-桌面应用开发文档-v2.docx"

FONT_CN = "Hiragino Sans GB"
FONT_CODE = "Menlo"

NAVY = "0B1F33"
BLUE = "006EFF"
BLUE_DARK = "0052D9"
INK = "17233D"
MUTED = "5E6D82"
GRID = "D9E2EC"
PALE_BLUE = "E8F3FF"
PALE_GRAY = "F4F6F9"
PALE_TEAL = "E8F8F5"
TEAL = "00A870"
AMBER = "ED7B2F"
RED = "D54941"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def style_run(run, size=11, color=INK, bold=False, italic=False, font=FONT_CN) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    style_run(run, size=9, color=MUTED)
    field_begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    field_begin_run._r.append(begin)
    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)
    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    field_text = paragraph.add_run("1")
    style_run(field_text, size=9, color=MUTED)
    field_end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_end_run._r.append(end)
    tail = paragraph.add_run(" 页")
    style_run(tail, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = FONT_CN
    title._element.rPr.rFonts.set(qn("w:ascii"), FONT_CN)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CN)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    title_p_pr = title._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    subtitle = styles["Subtitle"]
    subtitle.font.name = FONT_CN
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), FONT_CN)
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CN)
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.keep_with_next = True

    heading_tokens = {
        "Heading 1": (16, 18, 10, BLUE_DARK),
        "Heading 2": (13, 14, 7, BLUE_DARK),
        "Heading 3": (12, 10, 5, NAVY),
    }
    for style_name, (size, before, after, color) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", 1)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = FONT_CODE
    code_style._element.rPr.rFonts.set(qn("w:ascii"), FONT_CODE)
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CODE)
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CODE)
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string(NAVY)
    code_style.paragraph_format.left_indent = Inches(0.16)
    code_style.paragraph_format.right_indent = Inches(0.16)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(7)
    code_style.paragraph_format.line_spacing = 1.05


def configure_numbering(doc: Document) -> tuple[int, int, int]:
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abstract, default=0) + 1
    next_num = max(existing_num, default=0) + 1

    def add_definition(abstract_id: int, num_id: int, fmt: str, text_value: str, font: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), font)
        r_fonts.set(qn("w:hAnsi"), font)
        r_fonts.set(qn("w:eastAsia"), font)
        r_pr.append(r_fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)

    add_definition(next_abs, next_num, "bullet", "•", FONT_CN)
    add_definition(next_abs + 1, next_num + 1, "decimal", "%1.", FONT_CN)
    add_definition(next_abs + 2, next_num + 2, "decimal", "%1.", FONT_CN)
    return next_num, next_num + 1, next_num + 2


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color=GRID, size=6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), color)
        p_bdr.append(node)


def add_bullet(doc, text: str, bullet_num_id: int) -> None:
    p = doc.add_paragraph(style="List Bullet")
    apply_num(p, bullet_num_id)
    p.add_run(text)


def add_number(doc, text: str, number_num_id: int) -> None:
    p = doc.add_paragraph(style="List Number")
    apply_num(p, number_num_id)
    p.add_run(text)


def add_callout(doc, label: str, text: str, fill=PALE_BLUE, accent=BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, color=accent, size=8)
    lead = p.add_run(f"{label}  ")
    style_run(lead, size=10.5, color=accent, bold=True)
    body = p.add_run(text)
    style_run(body, size=10.5, color=INK)


def add_code_block(doc, code: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, "F7F9FC")
    set_paragraph_border(p, color=GRID, size=4)
    for idx, line in enumerate(code.splitlines()):
        run = p.add_run(line)
        style_run(run, size=8.5, color=NAVY, font=FONT_CODE)
        if idx != len(code.splitlines()) - 1:
            run.add_break()


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int], header_fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    table.paragraph_format if hasattr(table, "paragraph_format") else None
    header = table.rows[0]
    repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        style_run(r, size=9.5, color=NAVY, bold=True)
    for row_values in rows:
        row = table.add_row()
        keep_table_row_together(row)
        for idx, text in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.10
            if idx == 0 and len(headers) <= 3:
                run = p.add_run(text)
                style_run(run, size=9.3, color=INK, bold=True)
            else:
                run = p.add_run(text)
                style_run(run, size=9.3, color=INK)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_image_alt(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_section_intro(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    style_run(r, size=10.8, color=MUTED)


def page_break(doc) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_document() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if not MOCKUP_PATH.exists():
        raise FileNotFoundError(MOCKUP_PATH)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)
    bullet_num_id, number_num_id, acceptance_num_id = configure_numbering(doc)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    # Quiet running header/footer; first page uses the same compact furniture.
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    hr = header_p.add_run("玄壳  ·  桌面应用开发文档")
    style_run(hr, size=8.5, color=MUTED, bold=True)
    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)

    # Cover / memo masthead.
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    kr = kicker.add_run("PRODUCT & ENGINEERING SPECIFICATION")
    style_run(kr, size=9, color=BLUE, bold=True)
    title = doc.add_paragraph(style="Title")
    title.add_run("玄壳")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("C++ 桌面端服务器管理与实时运维控制台 · V1.0")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    for label, value in (
        ("文档定位", "产品需求 + 技术设计 + 验收基线"),
        ("目标平台", "Windows 10/11 优先，兼容 macOS 与主流 Linux 桌面"),
        ("服务器范围", "Linux 主机，SSH/SFTP 无代理接入"),
        ("生成日期", str(date.today())),
    ):
        lr = meta.add_run(f"{label}：")
        style_run(lr, size=9.5, color=NAVY, bold=True)
        vr = meta.add_run(f"{value}\n")
        style_run(vr, size=9.5, color=MUTED)

    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.space_after = Pt(4)
    shape = image_p.add_run().add_picture(str(MOCKUP_PATH), width=Inches(6.35))
    add_image_alt(
        shape,
        "玄壳桌面端效果图",
        "SSH 优先的紧凑型云运维桌面控制台，左侧仅保留主机导航，主工作区展示状态摘要、SSH 终端和 SFTP 文件管理。",
    )
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    cr = caption.add_run("图 1  ·  SSH 优先的高保真桌面端效果图（腾讯云式信息架构语言，不使用其品牌资产）")
    style_run(cr, size=8.5, color=MUTED, italic=True)

    add_callout(
        doc,
        "产品一句话",
        "面向开发、运维和小型技术团队的一体化 SSH 桌面工作台：在同一窗口中完成主机管理、秒级性能观测、交互式命令执行与服务器文件操作。",
    )

    doc.add_heading("1. 项目范围与设计原则", level=1)
    add_section_intro(doc, "本节将原始需求转化为可设计、可开发、可验收的产品边界。")

    doc.add_heading("1.1 需求解释与关键假设", level=2)
    add_bullet(doc, "“程序 C++”按原生 C++ 桌面应用理解，推荐 C++20 + Qt 6.x Widgets 技术栈。", bullet_num_id)
    add_bullet(doc, "“腾讯云布局风格”按紧凑云控制台的信息架构、深色侧栏与亮蓝强调色理解；不得复制商标、Logo 或受保护品牌素材。", bullet_num_id)
    add_bullet(doc, "原始需求中的“内核使用情况”按“内存使用情况”落地；扩展监控可另行加入内核态 CPU、上下文切换、进程/线程数。", bullet_num_id)
    add_bullet(doc, "V1 面向通过 SSH 可达的 Linux 服务器；Windows Server、容器编排和云厂商 API 接入列入后续版本。", bullet_num_id)

    doc.add_heading("1.2 产品目标", level=2)
    add_table(
        doc,
        ["目标", "V1 结果定义"],
        [
            ["统一入口", "服务器、监控、终端、文件四类工作流无需切换应用。"],
            ["实时可见", "选中主机后 2 秒内呈现首批指标，默认每 1 秒刷新。"],
            ["安全可控", "凭据不落明文；首次连接校验主机指纹；危险操作有确认与审计。"],
            ["桌面效率", "键盘优先、信息紧凑、多标签会话和可恢复的工作区状态。"],
        ],
        [1800, 7560],
    )

    doc.add_heading("1.3 V1 非目标", level=2)
    add_bullet(doc, "不替代 Prometheus、Zabbix 等长期时序监控平台；V1 以短期诊断和实时操作为主。", bullet_num_id)
    add_bullet(doc, "不内置公有云账户、计费、VPC 或工单管理；仅借鉴云控制台的布局语言。", bullet_num_id)
    add_bullet(doc, "不提供多人协作服务端与集中式堡垒机能力；桌面端本地运行、按用户隔离数据。", bullet_num_id)

    doc.add_heading("2. 用户、场景与信息架构", level=1)
    add_section_intro(doc, "界面围绕“定位主机 → 观察状态 → 执行命令/处理文件 → 留下审计记录”的闭环组织。")

    doc.add_heading("2.1 目标用户与主任务", level=2)
    add_table(
        doc,
        ["角色", "高频任务", "关键诉求"],
        [
            ["开发工程师", "排查服务状态、查看日志、部署文件", "连接快、终端顺手、文件传输可靠"],
            ["系统运维", "批量维护主机、观察负载、处理告警", "指标准确、多主机会话、操作可追溯"],
            ["技术负责人", "检查运行健康度、复核变更", "状态聚合、风险提示、最小权限"],
        ],
        [1700, 3800, 3860],
    )

    doc.add_heading("2.2 导航与窗口布局", level=2)
    add_table(
        doc,
        ["区域", "建议尺寸", "内容与行为"],
        [
            ["顶部栏", "48 px", "产品名、全局搜索、连接状态、账户菜单、窗口控制。"],
            ["主机导航", "230–244 px", "唯一左侧栏；仅放主机搜索、分组、状态点、名称/IP，底部保留连接与凭据入口；支持折叠。"],
            ["工作区标签", "40–44 px", "选中主机后切换监控、SSH 终端、文件管理；进入主机默认恢复最近使用标签，首次默认 SSH 终端。"],
            ["状态摘要", "110–130 px", "CPU、内存、负载、磁盘四项紧凑 KPI 与关键告警；完整趋势图只在“监控”标签展开。"],
            ["远程工作区", "自适应", "SSH 终端默认占宽约 62%，SFTP 文件管理约 38%；分隔条可拖拽，也可单独全屏。"],
            ["状态栏", "24 px", "采样延迟、传输队列、连接安全状态与版本信息。"],
        ],
        [1500, 1400, 6460],
    )
    add_callout(doc, "显示基线", "最小窗口 1280×720；推荐 1440×900；左侧始终只占一列。1920×1080 下完整显示状态摘要与终端/文件双工作区，终端获得主要面积。", fill=PALE_GRAY, accent=BLUE_DARK)

    doc.add_heading("2.3 视觉与交互规范", level=2)
    add_bullet(doc, "颜色：深海军蓝 #0B1F33、主强调蓝 #006EFF、成功绿 #00A870、警告橙 #ED7B2F、危险红 #D54941。", bullet_num_id)
    add_bullet(doc, "密度：基础间距 4/8 px；常规控件高 32 px；表格行高 32–36 px；卡片内边距 12–16 px。", bullet_num_id)
    add_bullet(doc, "导航：服务器是唯一主对象；监控、终端和文件是当前服务器的工作区标签，不在左侧重复形成一级导航。", bullet_num_id)
    add_bullet(doc, "状态不仅依赖颜色：同时使用图标、文字与数值；图表须有单位、图例、时间范围和空数据态。", bullet_num_id)
    add_bullet(doc, "键盘：Ctrl/Cmd+K 全局搜索，Ctrl/Cmd+T 新建终端，Ctrl/Cmd+L 聚焦路径，F5 刷新。", bullet_num_id)

    doc.add_heading("3. 功能需求", level=1)
    add_section_intro(doc, "以下条目是产品需求与开发拆分的共同基线；优先级采用 P0（必须）、P1（应有）、P2（可延后）。")

    doc.add_heading("3.1 服务器管理", level=2)
    add_table(
        doc,
        ["ID", "优先级", "需求与验收要点"],
        [
            ["SRV-01", "P0", "新增、编辑、复制、删除服务器；字段含名称、主机/IP、端口、用户名、分组、标签和备注。"],
            ["SRV-02", "P0", "支持密码、私钥、SSH Agent 三种认证；敏感值仅写入系统凭据库，SQLite 只保存引用。"],
            ["SRV-03", "P0", "“测试连接”返回 DNS、TCP、SSH 握手、认证、SFTP 子系统的分阶段结果。"],
            ["SRV-04", "P0", "首次连接展示主机指纹；用户确认后写入 known_hosts；指纹变化必须阻断并提示。"],
            ["SRV-05", "P1", "支持跳板机、KeepAlive、连接超时和字符编码设置；列表可按分组/状态/标签过滤。"],
            ["SRV-06", "P1", "导入/导出非敏感配置；导出文件不得包含密码、私钥内容或令牌。"],
        ],
        [1200, 900, 7260],
    )

    doc.add_heading("3.2 实时观测", level=2)
    add_table(
        doc,
        ["指标", "采集口径（Linux）", "展示与告警"],
        [
            ["CPU", "/proc/stat 两次采样差分，区分 user/system/iowait；默认 1 秒间隔。", "总使用率 KPI + 60 分钟折线；>85% 持续 5 分钟警告。"],
            ["内存", "/proc/meminfo；used = total - available，兼容无 MemAvailable 的旧内核。", "已用/总量、百分比、Swap；>90% 警告。"],
            ["负载", "/proc/loadavg 的 1/5/15 分钟值，并按逻辑 CPU 核数归一化。", "三条曲线；1 分钟归一负载 >1.0 警告。"],
            ["磁盘", "优先解析 df -P -B1；排除 tmpfs、devtmpfs，可配置挂载点。", "容量、已用、剩余、挂载点；>80% 警告、>90% 严重。"],
            ["扩展指标", "网络收发、进程数、上下文切换、内核态 CPU、文件句柄。", "P2 插件式指标卡，不阻塞 V1。"],
        ],
        [1600, 4300, 3460],
    )
    add_bullet(doc, "连接后 2 秒内首屏有数据；断线保留最后值并标记“数据已过期”，不得显示为正常实时值。", bullet_num_id)
    add_bullet(doc, "内存环形缓冲默认保留 60 分钟；可选写入 SQLite，默认保留 7 天并按天清理。", bullet_num_id)
    add_bullet(doc, "支持 1/5/15/30 秒采样间隔；当应用进入后台或主机未选中时自动降频。", bullet_num_id)

    doc.add_heading("3.3 SSH 命令行", level=2)
    add_table(
        doc,
        ["ID", "优先级", "需求与验收要点"],
        [
            ["TERM-01", "P0", "打开交互式 shell channel 并请求 PTY（xterm-256color）；窗口缩放同步终端行列。"],
            ["TERM-02", "P0", "支持 ANSI/VT100 常用序列、UTF-8、复制粘贴、滚动回看、搜索和清屏。"],
            ["TERM-03", "P0", "多标签会话；标签显示主机名、连接状态和未读输出；异常断线可一键重连。"],
            ["TERM-04", "P1", "多行或超过 200 字符的粘贴默认二次确认，降低误执行风险。"],
            ["TERM-05", "P1", "命令历史默认仅保存在内存；如开启本地历史，支持禁用与按主机清除。"],
        ],
        [1200, 900, 7260],
    )

    doc.add_heading("3.4 服务器文件管理", level=2)
    add_table(
        doc,
        ["ID", "优先级", "需求与验收要点"],
        [
            ["FILE-01", "P0", "通过 SFTP 列出目录，展示名称、大小、修改时间、权限、所有者和类型；支持面包屑与手输路径。"],
            ["FILE-02", "P0", "上传、下载、新建文件夹、重命名、删除、刷新；删除和覆盖必须确认。"],
            ["FILE-03", "P0", "大文件流式传输，显示速度、进度、剩余时间；支持取消、失败重试和断点续传（服务端允许时）。"],
            ["FILE-04", "P1", "上传先写入 .part 临时文件，完成校验后原子重命名，避免半成品覆盖目标。"],
            ["FILE-05", "P1", "目录 10,000 项仍可滚动；分页/虚拟化渲染，排序不阻塞 UI。"],
            ["FILE-06", "P1", "冲突策略：询问、覆盖、跳过、重命名；批量任务支持“全部应用”。"],
        ],
        [1200, 900, 7260],
    )

    doc.add_heading("3.5 任务中心、通知与审计", level=2)
    add_bullet(doc, "任务中心统一展示连接、采样、文件传输、批量操作的状态、错误与重试入口。", bullet_num_id)
    add_bullet(doc, "告警支持阈值、持续时长和恢复事件；桌面通知默认对重复事件做 5 分钟抑制。", bullet_num_id)
    add_bullet(doc, "审计记录操作人、本地主机、目标服务器、动作、结果、耗时和时间戳；默认不记录密码、私钥与完整终端内容。", bullet_num_id)

    doc.add_heading("4. 技术方案与总体架构", level=1)
    add_section_intro(doc, "采用模块化单体桌面架构：UI、应用服务与基础设施解耦，以接口隔离便于后续替换 SSH 库或引入监控代理。")

    doc.add_heading("4.1 推荐技术栈", level=2)
    add_table(
        doc,
        ["层面", "选型", "说明"],
        [
            ["语言与构建", "C++20、CMake 3.25+、Ninja", "统一 Windows/macOS/Linux 构建；启用 clang-tidy、warnings-as-errors（CI）。"],
            ["桌面 UI", "Qt 6.x Widgets", "适合紧凑型多面板桌面应用；Model/View 处理大列表，QSplitter 组织工作区。"],
            ["SSH/SFTP", "libssh2（非阻塞模式）", "BSD 许可；同一连接复用 shell、exec、SFTP channel，外层封装连接池。"],
            ["终端渲染", "自研 TerminalWidget + libvterm 类解析器", "隔离 VT 状态机与渲染层；避免将 shell 输出当富文本直接追加。"],
            ["本地数据", "SQLite + 迁移脚本", "保存服务器元数据、偏好、短期指标与审计；敏感凭据走系统凭据库。"],
            ["日志与测试", "spdlog、Catch2/GoogleTest、Qt Test", "结构化日志、单元/集成/UI 冒烟测试；依赖版本锁定。"],
            ["图表", "自定义 QPainter 轻量曲线", "环形缓冲 + 降采样；如采用 Qt Charts，需先确认许可证与发布模式。"],
        ],
        [1600, 3000, 4760],
    )

    doc.add_heading("4.2 模块边界", level=2)
    add_table(
        doc,
        ["层", "核心组件", "职责"],
        [
            ["Presentation", "MainWindow、ServerListModel、MonitorView、TerminalTabs、FileBrowser", "布局、输入校验、状态呈现；不得直接调用 libssh2。"],
            ["Application", "ServerService、MonitorService、TerminalService、FileService、TaskService", "编排用例、权限判断、取消/重试、向 UI 发布不可变状态。"],
            ["Domain", "ServerProfile、MetricSample、TransferTask、AlertRule、AuditEvent", "业务模型、状态机、阈值与错误语义；不依赖 Qt Widgets。"],
            ["Infrastructure", "SshSession、SftpClient、LinuxProbe、SqliteRepository、SecretStore", "网络 I/O、解析、持久化、系统凭据库和平台适配。"],
        ],
        [1800, 3700, 3860],
    )

    doc.add_heading("4.3 线程与并发模型", level=2)
    add_number(doc, "UI 主线程仅负责渲染、Model/View 更新和用户输入；任何 DNS、SSH、磁盘或数据库操作不得阻塞 UI。", number_num_id)
    add_number(doc, "每个活跃主机由一个 Session Actor 串行管理 SSH 状态，避免同一 libssh2 session 被跨线程并发调用。", number_num_id)
    add_number(doc, "采样器使用单独调度线程和有界任务队列；同主机的指标命令合并为一次 exec，结果一次解析。", number_num_id)
    add_number(doc, "终端与 SFTP 使用独立 channel；文件传输按主机默认并发 2 个，全局默认 4 个，支持取消令牌。", number_num_id)
    add_number(doc, "跨线程消息只传值对象；高频指标采用 ring buffer，UI 每秒最多重绘一次并按像素宽度降采样。", number_num_id)

    doc.add_heading("4.4 核心接口草案", level=2)
    add_code_block(
        doc,
        "struct MetricSample {\n"
        "    std::chrono::system_clock::time_point ts;\n"
        "    double cpuPercent{}, memoryPercent{};\n"
        "    double load1{}, load5{}, load15{};\n"
        "    std::vector<DiskUsage> disks;\n"
        "};\n\n"
        "class IRemoteSession {\n"
        "public:\n"
        "    virtual Task<Result<void>> connect(const ServerProfile&, CancelToken) = 0;\n"
        "    virtual Task<Result<ExecResult>> exec(std::string_view, CancelToken) = 0;\n"
        "    virtual Task<Result<ShellHandle>> openShell(const PtyOptions&, CancelToken) = 0;\n"
        "    virtual Task<Result<std::shared_ptr<ISftpClient>>> openSftp(CancelToken) = 0;\n"
        "    virtual ~IRemoteSession() = default;\n"
        "};"
    )

    doc.add_heading("5. 数据设计与采集协议", level=1)
    add_section_intro(doc, "本地数据按“配置、短时状态、任务、审计”分层；凭据与业务数据物理分离。")

    doc.add_heading("5.1 核心数据对象", level=2)
    add_table(
        doc,
        ["对象", "关键字段", "持久化策略"],
        [
            ["ServerProfile", "id、name、host、port、user、group、tags、jumpHostId、credentialRef", "SQLite；不含秘密内容"],
            ["CredentialRef", "provider、key、authType", "SQLite 仅保存引用；秘密在 Credential Manager/Keychain/Secret Service"],
            ["MetricSample", "serverId、timestamp、cpu、memory、load、disks", "内存环形缓冲；可选 SQLite 7 天"],
            ["TerminalSession", "id、serverId、createdAt、state、ptySize", "工作区元数据可存；输出默认不落盘"],
            ["TransferTask", "source、target、direction、bytes、state、retryCount", "任务完成后保留 30 天"],
            ["AuditEvent", "actor、serverId、action、result、duration、createdAt", "SQLite 90 天，可配置导出"],
        ],
        [2100, 4600, 2660],
    )

    doc.add_heading("5.2 Linux 无代理采集", level=2)
    add_code_block(
        doc,
        "LC_ALL=C sh -c '\n"
        "  printf \"__CPU__\\n\";  head -n 1 /proc/stat;\n"
        "  printf \"__MEM__\\n\";  cat /proc/meminfo;\n"
        "  printf \"__LOAD__\\n\"; cat /proc/loadavg;\n"
        "  printf \"__DISK__\\n\"; df -P -B1 -x tmpfs -x devtmpfs\n"
        "'"
    )
    add_bullet(doc, "命令固定 LC_ALL=C，解析器按分隔标记切片，拒绝依赖本地化列名。", bullet_num_id)
    add_bullet(doc, "每次采集设置 3 秒超时和 1 MiB 输出上限；超限、超时或部分缺失均产生明确质量标记。", bullet_num_id)
    add_bullet(doc, "服务器时间只用于展示诊断；所有样本索引使用本机单调时钟，落库使用 UTC 时间戳。", bullet_num_id)

    doc.add_heading("5.3 连接状态机", level=2)
    add_table(
        doc,
        ["状态", "进入条件", "允许动作 / 退出条件"],
        [
            ["Disconnected", "初始、用户断开或不可恢复错误", "Connect；成功后进入 Authenticating"],
            ["Connecting", "DNS/TCP 建连中", "Cancel；TCP 成功进入 VerifyingHostKey"],
            ["VerifyingHostKey", "收到主机公钥", "已知且匹配继续；未知需确认；变化则 Failed"],
            ["Authenticating", "主机指纹通过", "认证成功进入 Ready；失败可重试或换凭据"],
            ["Ready", "SSH 会话可用", "打开监控/终端/SFTP channel；网络异常进入 Reconnecting"],
            ["Reconnecting", "意外断线且策略允许", "指数退避 1/2/4/8/15 秒；成功回 Ready；上限后 Failed"],
            ["Failed", "认证、指纹、协议或重试耗尽", "展示阶段化错误；修改配置后 Connect"],
        ],
        [2400, 2700, 4260],
    )

    doc.add_heading("6. 安全、权限与审计", level=1)
    add_section_intro(doc, "默认安全策略必须在首次使用时成立，而不是依赖用户进入设置页主动开启。")
    add_table(
        doc,
        ["控制项", "V1 要求"],
        [
            ["凭据保护", "密码、私钥口令写入操作系统凭据库；内存中的敏感缓冲使用后清零；日志自动脱敏。"],
            ["主机身份", "严格 known_hosts 校验；未知指纹需人工确认；指纹变化阻断，禁止默认 Auto-accept。"],
            ["最小权限", "推荐使用专用运维账户；不自动请求 root；需要 sudo 的动作必须由用户显式执行。"],
            ["终端安全", "长/多行粘贴确认；可配置禁止保存历史；审计默认只记录会话元数据，不记录敏感输出。"],
            ["文件安全", "路径规范化，禁止 ../ 越界；覆盖/删除确认；上传临时文件校验后原子重命名。"],
            ["本地数据", "SQLite 使用受限文件权限；提供“一键清除本地数据”和审计导出；崩溃报告不包含秘密。"],
            ["供应链", "依赖锁定、SBOM、签名安装包、自动更新签名校验；CI 运行 SAST 与依赖漏洞扫描。"],
        ],
        [1900, 7460],
    )

    doc.add_heading("7. 错误处理与可观测性", level=1)
    add_table(
        doc,
        ["错误类别", "用户提示", "内部处理"],
        [
            ["网络不可达", "“无法连接到 host:port”，展示 DNS/TCP 阶段和重试入口。", "指数退避；记录 errno/平台错误码；不循环弹窗。"],
            ["认证失败", "区分用户名、密码/密钥、Agent、权限拒绝。", "清理会话；限制自动重试；不记录凭据。"],
            ["主机指纹变化", "高风险阻断页，展示旧/新指纹与处理建议。", "不自动覆盖 known_hosts；写入安全审计。"],
            ["采集解析失败", "卡片显示“数据不可用/已过期”，其他模块仍可操作。", "保存原始片段摘要与解析器版本，避免输出敏感信息。"],
            ["SFTP 失败", "任务行展示路径、阶段和可重试原因。", "保留 .part 或按策略清理；校验已写入字节。"],
            ["应用崩溃", "下次启动提示恢复工作区，不自动重连敏感会话。", "本地 minidump 可选；日志轮转 10×10 MiB。"],
        ],
        [1800, 3600, 3960],
    )

    doc.add_heading("8. 性能与质量指标", level=1)
    add_table(
        doc,
        ["指标", "目标值", "测量方法"],
        [
            ["冷启动", "P95 ≤ 2.5 秒", "标准办公电脑、100 台主机配置、首次主窗体可交互"],
            ["交互响应", "P95 ≤ 100 ms", "导航、筛选、切换标签不含网络等待"],
            ["监控首屏", "≤ 2 秒", "SSH 已连接后出现第一批完整 KPI"],
            ["监控容量", "50 台已保存、10 台 1 秒活跃采样", "持续 8 小时，UI 无明显卡顿，丢样率 <0.1%"],
            ["资源占用", "空闲 CPU <1%；10 台采样 CPU <5%；内存 <300 MiB", "Release 构建，排除文件高速传输"],
            ["文件列表", "10,000 项首屏 ≤ 1 秒；滚动 60 fps 目标", "本地模拟 SFTP 与真实高延迟环境各测试一次"],
            ["可靠性", "24 小时稳定运行，无持续增长泄漏", "ASan/LSan、长稳测试、断网/恢复故障注入"],
        ],
        [1900, 2500, 4960],
    )

    doc.add_heading("9. 测试策略", level=1)
    doc.add_heading("9.1 自动化测试", level=2)
    add_bullet(doc, "单元测试：/proc 与 df 解析器、状态机、阈值/恢复逻辑、路径规范化、配置迁移。", bullet_num_id)
    add_bullet(doc, "集成测试：Docker/OpenSSH 测试靶机覆盖密码、密钥、错误凭据、指纹变化、SFTP 权限和网络延迟。", bullet_num_id)
    add_bullet(doc, "UI 测试：服务器 CRUD、切换主机、终端标签、拖拽分隔条、文件冲突对话框和键盘快捷键。", bullet_num_id)
    add_bullet(doc, "安全测试：日志脱敏、known_hosts 阻断、目录穿越、符号链接、超长路径、恶意 ANSI 序列。", bullet_num_id)
    add_bullet(doc, "跨平台冒烟：Windows 10/11、macOS 当前两个主版本、Ubuntu LTS；高 DPI 100/150/200%。", bullet_num_id)

    doc.add_heading("9.2 核心验收场景", level=2)
    add_number(doc, "新增一台使用私钥认证的 Linux 主机，确认指纹后连接成功；重启应用后凭据仍可用且数据库中无明文。", acceptance_num_id)
    add_number(doc, "选中主机后 2 秒内出现 CPU、内存、负载与磁盘；拔网线后显示已过期，恢复网络后自动重连并继续曲线。", acceptance_num_id)
    add_number(doc, "打开两个终端标签执行 top/htop、中文输出和 ANSI 彩色命令；缩放窗口后行列正确且 UI 无阻塞。", acceptance_num_id)
    add_number(doc, "向 /var/tmp 上传 1 GiB 文件，暂停网络再恢复；任务可重试/续传，目标无半文件覆盖。", acceptance_num_id)
    add_number(doc, "浏览含 10,000 项的目录，筛选、排序和滚动保持可用；无权限目录给出明确错误而非空白页。", acceptance_num_id)

    doc.add_heading("10. 研发计划与交付物", level=1)
    add_section_intro(doc, "以下按 2 名前端/桌面 C++ + 1 名测试的典型小团队估算，可根据人员并行度压缩。")
    add_table(
        doc,
        ["阶段", "周期", "主要交付"],
        [
            ["M1 基础工程", "第 1–2 周", "CMake/CI、Qt 主框架、主题与布局、SQLite 迁移、日志、设置。"],
            ["M2 服务器与 SSH", "第 3–4 周", "服务器 CRUD、凭据库、指纹校验、连接状态机、连接测试。"],
            ["M3 实时监控", "第 5–6 周", "LinuxProbe、指标模型、KPI/曲线、断线与过期状态、基础告警。"],
            ["M4 交互终端", "第 7–8 周", "PTY、VT 解析/渲染、多标签、粘贴保护、重连。"],
            ["M5 文件管理", "第 9–10 周", "SFTP 浏览、传输任务、冲突策略、取消/重试、目录大数据优化。"],
            ["M6 稳定与发布", "第 11–12 周", "安全/性能/长稳测试、安装包签名、更新通道、用户手册。"],
        ],
        [1800, 1400, 6160],
    )

    doc.add_heading("11. Definition of Done", level=1)
    add_bullet(doc, "P0 需求全部通过自动化或可复现手工验收；P1 未完成项有明确版本计划。", bullet_num_id)
    add_bullet(doc, "Windows Release 安装包可签名安装/卸载；配置升级可回滚；无管理员权限也可运行。", bullet_num_id)
    add_bullet(doc, "CI 通过编译、单元/集成测试、静态分析、依赖漏洞扫描和许可证清单生成。", bullet_num_id)
    add_bullet(doc, "24 小时长稳、断网恢复、10 台秒级采样、1 GiB 文件传输和 10,000 项目录场景通过。", bullet_num_id)
    add_bullet(doc, "安全评审确认：凭据不落明文、指纹变化阻断、日志脱敏、危险文件操作可审计。", bullet_num_id)
    add_bullet(doc, "交付源码、构建说明、架构说明、用户手册、测试报告、SBOM、安装包与版本发布说明。", bullet_num_id)

    doc.add_heading("12. 后续演进", level=1)
    add_table(
        doc,
        ["版本", "方向", "候选能力"],
        [
            ["V1.1", "效率增强", "批量命令、命令片段、文件对比、会话布局模板、告警静默。"],
            ["V1.2", "平台扩展", "Windows Server/PowerShell、代理式高频指标、系统服务与进程管理。"],
            ["V2.0", "团队协作", "集中策略、角色权限、共享主机目录、集中审计、堡垒机/零信任接入。"],
            ["V2.x", "生态集成", "Prometheus、云厂商 API、Kubernetes、工单/IM 通知与插件 SDK。"],
        ],
        [1500, 1800, 6060],
    )

    # Core properties and save.
    props = doc.core_properties
    props.title = "玄壳 - C++ 桌面应用开发文档"
    props.subject = "服务器管理、实时监控、SSH 终端与 SFTP 文件管理"
    props.author = "Product & Engineering"
    props.keywords = "C++, Qt, SSH, SFTP, 服务器管理, 实时监控, 桌面应用"
    props.comments = "Generated development specification"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
